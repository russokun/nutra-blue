import io
import re
import requests
import unicodedata
from docx import Document

def _normalize_text(text: str) -> str:
    """Quita acentos y pasa a minúsculas para comparaciones robustas."""
    text = (text or "").strip().lower()
    return "".join(c for c in unicodedata.normalize("NFD", text) if unicodedata.category(c) != "Mn")

def parse_google_doc(url: str) -> dict:
    """
    Descarga un Google Doc como archivo .docx y extrae la información estructurada
    de las secciones de producto:
    - description (Qué es / Descripción del producto)
    - product_profile (Qué lo hace distinto / Sabor y textura / Perfil)
    - origin (Origen / Zona de producción)
    - cross_selling (Va bien con / Recomendaciones de venta cruzada)
    - ingredients (Ingredientes / Composición)
    - usage (Cómo se toma / Cómo se come / Dosis / Porción)
    - precautions (Antes de consumir / Advertencias / Contraindicaciones)
    - extracted_benefits (Para qué te sirve / Viñetas de beneficios clave)
    """
    # Limpiar y validar URL
    url = url.strip() if url else ""
    if not url:
        return {}

    # Extraer el ID del documento
    match = re.search(r"/document/d/([a-zA-Z0-9-_]+)", url)
    if not match:
        raise ValueError(f"URL de Google Doc no válida: {url}")
    
    doc_id = match.group(1)
    export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=docx"
    
    # Descargar el documento como binario
    try:
        response = requests.get(export_url, timeout=15)
        response.raise_for_status()
    except Exception as e:
        raise RuntimeError(f"No se pudo descargar el documento desde Google: {str(e)}")
    
    # Parsear el documento docx
    try:
        doc = Document(io.BytesIO(response.content))
    except Exception as e:
        raise RuntimeError(f"Error al abrir el archivo docx: {str(e)}")
    
    # Mapeo de secciones
    sections = {
        "description": "",
        "type_description": "",
        "origin": "",
        "cross_selling": "",
        "product_profile": "",
        "ingredients": "",
        "usage": "",
        "precautions": "",
        "flavor_profile": "",
        "extracted_benefits": []
    }

    current_key = None
    intro_captured = False

    def detect_header(raw_line: str) -> str:
        line_clean = raw_line.strip()
        norm = _normalize_text(line_clean)
        norm_no_bullets = re.sub(r"^[0-9\.\-\*\#\:\•\–\—\s]+", "", norm).strip()
        
        is_short = len(line_clean) <= 45 or line_clean.endswith(":")

        if not is_short and not any(norm_no_bullets.startswith(prefix) for prefix in [
            "va bien con:", "combina con:", "ingredientes:", "modo de uso:", "como se toma:", "como se come:"
        ]):
            return None

        # 1. Descripción / Qué es
        if norm_no_bullets in ["que es", "descripcion", "descripcion del producto", "sobre este producto"] or norm_no_bullets.startswith("que es:"):
            return "description"
        
        # 2. Beneficios / Para qué te sirve
        if norm_no_bullets in ["para que te sirve", "para que sirve", "beneficios para el cliente", "beneficios", "beneficios clave"]:
            return "benefits"
            
        # 3. Perfil del producto / Qué lo hace distinto
        if norm_no_bullets in ["que lo hace distinto", "que la hace distinta", "por que es especial", "perfil del producto", "perfil producto", "diferenciadores", "caracteristicas"]:
            return "product_profile"
            
        # 4. Sabor / Sensorial
        if norm_no_bullets in ["a que sabe", "sabor", "perfil sensorial", "textura", "sabor y textura"]:
            return "flavor_profile"
            
        # 5. Modo de uso / Cómo se toma / Cómo se come
        if norm_no_bullets in ["como se toma", "como se come", "como se usa", "modo de uso", "forma de uso", "modo de empleo", "instrucciones", "dosis", "cuanto comer", "cuanto tomar"]:
            return "usage"
            
        # 6. Venta cruzada / Combinaciones
        if any(norm_no_bullets.startswith(k) for k in ["va bien con", "combina con", "recomendaciones con otros productos", "cross-selling", "cross selling", "venta cruzada", "ideal con"]):
            return "cross_selling"
            
        # 7. Precauciones / Advertencias
        if norm_no_bullets in ["antes de consumir", "precauciones", "advertencias", "contraindicaciones", "cuidados", "importante"]:
            return "precautions"
            
        # 8. Origen / Zona de producción
        if norm_no_bullets in ["zona de produccion", "lugar de produccion", "pais de origen", "origen"]:
            return "origin"
            
        # 9. Ingredientes / Composición
        if norm_no_bullets in ["ingredientes", "ingrediente", "composicion", "formula"]:
            return "ingredients"
            
        # 10. Metadatos de costos o notas internas / cierre de página
        if any(h in norm_no_bullets for h in ["estructura de costos", "costos y margenes", "margenes", "estrategia de venta", "contacto de proveedor", "cada beneficio de esta hoja indica"]):
            return "ignored_metadata"

        return None

    # 1. Procesar párrafos
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
            
        # Al llegar al aviso legal de cierre de página, ignorar el pie de página
        if "cada beneficio de esta hoja" in _normalize_text(text):
            current_key = "ignored_metadata"
            continue

        header_key = detect_header(text)

        if header_key:
            current_key = header_key
            if ":" in text:
                resto = text.split(":", 1)[1].strip()
                if resto and current_key not in ("benefits", "ignored_metadata"):
                    sections[current_key] += resto + "\n"
            continue

        if current_key == "benefits":
            sub_lines = re.split(r"[\n\r\u000b]+", text)
            for line in sub_lines:
                clean = re.sub(r"^[^a-zA-Z0-9\s]{1,3}\s*", "", line.strip()).strip()
                if clean:
                    # Omitir etiquetas de validación técnica
                    if re.match(r"^respaldo\s+(solido|bueno|inicial|parcial)", _normalize_text(clean)):
                        continue
                    sections["extracted_benefits"].append(clean)
            continue

        if current_key is None:
            # Capturar intro o descripción inicial
            if not intro_captured and len(text) <= 60 and not any(k in _normalize_text(text) for k in ["superalimento", "fruto seco", "extracto"]):
                continue
            if "ficha comercial:" in text.lower() or text.lower().startswith("ficha"):
                continue
            sections["description"] += text + "\n"
            intro_captured = True
            continue

        if current_key and current_key != "ignored_metadata":
            sections[current_key] += text + "\n"

    # 2. Procesar tablas (crucial para 'Antes de consumir' / 'Cuánto comer' / advertencias / dosis)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if not cell_text:
                    continue
                
                cell_lines = [l.strip() for l in cell_text.splitlines() if l.strip()]
                for line in cell_lines:
                    line_header = detect_header(line)
                    if line_header:
                        table_key = line_header
                        if ":" in line:
                            resto = line.split(":", 1)[1].strip()
                            if resto and table_key not in ("benefits", "ignored_metadata"):
                                sections[table_key] += resto + "\n"
                        continue
                    
                    norm_line = _normalize_text(line)
                    if any(w in norm_line for w in [
                        "alergico", "alergia", "atragantamiento", "anticoagulante", "embarazada", 
                        "calculos renales", "menores de", "no la consumas", "no apta", "consulte a su medico", "consulta con tu medico"
                    ]):
                        clean_p = re.sub(r"^[^a-zA-Z0-9\s]{1,3}\s*", "", line).strip()
                        if clean_p and clean_p not in sections["precautions"]:
                            sections["precautions"] += ("• " if not clean_p.startswith("•") else "") + clean_p + "\n"
                    elif any(w in norm_line for w in [
                        "cuanto comer", "cuanto tomar", "porcion", "dientes al dia", "cucharadas al dia", "al dia", "por porcion"
                    ]):
                        clean_u = re.sub(r"^[^a-zA-Z0-9\s]{1,3}\s*", "", line).strip()
                        if clean_u and clean_u not in sections["usage"]:
                            sections["usage"] += clean_u + "\n"

    # 3. Deducir origen si no se capturó explícitamente pero se menciona en el perfil o descripción
    if not sections["origin"]:
        full_text = f"{sections['description']} {sections['product_profile']}"
        if re.search(r"\b(chiloe|chiloé|la araucania|la araucanía)\b", full_text, re.I):
            sections["origin"] = "Chile (Chiloé y La Araucanía)"
        elif re.search(r"\b(origen\s+chile|cultivada?\s+en\s+chile|hecho\s+en\s+chile|almendra\s+chilena|cosecha\s+en\s+chile|ciruela\s+chilena|avellana\s+chilena)\b", full_text, re.I):
            sections["origin"] = "Chile"
        elif re.search(r"\b(origen\s+china|producto\s+de\s+origen\s+china)\b", full_text, re.I):
            sections["origin"] = "China"
        elif re.search(r"\b(origen\s+peru|origen\s+perú|peru|perú)\b", full_text, re.I):
            sections["origin"] = "Perú"

    # 4. Si sabor (flavor_profile) se capturó, enriquecer el perfil del producto
    if sections["flavor_profile"]:
        sabor_txt = f"Sabor y textura: {sections['flavor_profile']}".strip()
        if sections["product_profile"]:
            sections["product_profile"] = f"{sections['product_profile']}\n{sabor_txt}".strip()
        else:
            sections["product_profile"] = sabor_txt

    # 5. Deducir ingredientes si está vacío pero la descripción o perfil indica pureza
    if not sections["ingredients"]:
        full_text = f"{sections['description']} {sections['product_profile']}"
        if any(term in full_text.lower() for term in ["naturales: crudas", "sin sal ni aceite", "sin conservantes", "100%", "naturales"]):
            sections["ingredients"] = "100% natural, sin aditivos, sal ni preservantes añadidos."

    # Limpiar espacios en blanco al inicio/final de cada sección
    cleaned_sections = {k: v.strip() if isinstance(v, str) else v for k, v in sections.items()}
    return cleaned_sections
